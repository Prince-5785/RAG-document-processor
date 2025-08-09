"""
Document processing module for handling various file formats including PDFs, Word documents, and emails.
"""

import os
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# PDF processing
try:
    import pypdf
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing libraries not available. Install pypdf and pdfplumber.")

# Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("Word document processing not available. Install python-docx.")

from .utils import clean_text, get_file_hash, validate_file_type

# Text chunking
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    logging.warning("LangChain not available. Text chunking will use basic splitting.")


class DocumentProcessor:
    """Handles document loading and text extraction from various file formats."""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.supported_formats = config.get('ui', {}).get('supported_formats', ['pdf', 'docx', 'eml', 'txt'])
        self.logger = logging.getLogger(__name__)
        
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a single file and extract text content with metadata.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Dictionary containing extracted text, metadata, and processing info
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if not validate_file_type(file_path, self.supported_formats):
            raise ValueError(f"Unsupported file format: {Path(file_path).suffix}")
            
        file_extension = Path(file_path).suffix.lower().lstrip('.')
        file_hash = get_file_hash(file_path)
        
        self.logger.info(f"Processing file: {file_path}")
        
        try:
            if file_extension == 'pdf':
                content = self._process_pdf(file_path)
            elif file_extension == 'docx':
                content = self._process_docx(file_path)
            elif file_extension == 'eml':
                content = self._process_email(file_path)
            elif file_extension == 'txt':
                content = self._process_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
            return {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_hash': file_hash,
                'file_type': file_extension,
                'content': content,
                'metadata': {
                    'file_size': os.path.getsize(file_path),
                    'processed_successfully': True
                }
            }
            
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {str(e)}")
            return {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_hash': file_hash,
                'file_type': file_extension,
                'content': "",
                'metadata': {
                    'file_size': os.path.getsize(file_path),
                    'processed_successfully': False,
                    'error': str(e)
                }
            }
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
            
        text_content = []
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"[Page {page_num}]\n{text}")
                        
        except Exception as e:
            self.logger.warning(f"pdfplumber failed for {file_path}, trying pypdf: {e}")

            # Fallback to pypdf
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            text_content.append(f"[Page {page_num}]\n{text}")
            except Exception as e2:
                self.logger.error(f"Both PDF processors failed for {file_path}: {e2}")
                raise e2
        
        return clean_text('\n\n'.join(text_content))
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("Word document processing not available")
            
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return clean_text('\n\n'.join(text_content))
    
    def _process_email(self, file_path: str) -> str:
        """Extract text from email file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            msg = email.message_from_file(file)
        
        text_content = []
        
        # Extract headers
        subject = msg.get('Subject', '')
        sender = msg.get('From', '')
        recipient = msg.get('To', '')
        date = msg.get('Date', '')
        
        if subject:
            text_content.append(f"Subject: {subject}")
        if sender:
            text_content.append(f"From: {sender}")
        if recipient:
            text_content.append(f"To: {recipient}")
        if date:
            text_content.append(f"Date: {date}")
        
        text_content.append("")  # Empty line separator
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    body = part.get_payload(decode=True)
                    if body:
                        text_content.append(body.decode('utf-8', errors='ignore'))
        else:
            body = msg.get_payload(decode=True)
            if body:
                text_content.append(body.decode('utf-8', errors='ignore'))
        
        return clean_text('\n\n'.join(text_content))
    
    def _process_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
        
        return clean_text(content)
    
    def process_multiple_files(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Process multiple files and return their extracted content.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            List of dictionaries containing processed file information
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = self.process_file(file_path)
                results.append(result)
            except Exception as e:
                self.logger.error(f"Failed to process file {file_path}: {e}")
                results.append({
                    'file_path': file_path,
                    'file_name': Path(file_path).name,
                    'content': "",
                    'metadata': {
                        'processed_successfully': False,
                        'error': str(e)
                    }
                })
        
        return results
    
    def get_supported_formats(self) -> List[str]:
        """Return list of supported file formats."""
        return self.supported_formats.copy()

    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks for embedding.

        Args:
            text: Text content to chunk
            metadata: Optional metadata to include with each chunk

        Returns:
            List of dictionaries containing chunk text and metadata
        """
        if not text.strip():
            return []

        doc_config = self.config.get('document_processing', {})
        chunk_size = doc_config.get('chunk_size', 1000)
        chunk_overlap = doc_config.get('chunk_overlap', 200)
        separators = doc_config.get('separators', ["\n\n", "\n", " ", ""])

        if LANGCHAIN_AVAILABLE:
            # Use LangChain's RecursiveCharacterTextSplitter
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=separators,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
        else:
            # Fallback to basic chunking
            chunks = self._basic_text_chunking(text, chunk_size, chunk_overlap)

        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                'chunk_index': i,
                'chunk_size': len(chunk),
                'total_chunks': len(chunks)
            }

            if metadata:
                chunk_metadata.update(metadata)

            chunk_objects.append({
                'text': chunk,
                'metadata': chunk_metadata
            })

        return chunk_objects

    def _basic_text_chunking(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """
        Basic text chunking implementation when LangChain is not available.

        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            chunk_overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            if end >= len(text):
                chunks.append(text[start:])
                break

            # Try to break at sentence boundary
            chunk = text[start:end]
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')

            # Find the best break point
            break_point = max(last_period, last_newline)
            if break_point > start + chunk_size // 2:  # Only if break point is reasonable
                end = start + break_point + 1

            chunks.append(text[start:end])
            start = end - chunk_overlap

        return chunks
